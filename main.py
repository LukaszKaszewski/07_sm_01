# czesc druga

import numpy as np
import matplotlib.pyplot as plt
import sounddevice as sd
import soundfile as sf
import scipy.fftpack
from docx import Document
from docx.shared import Inches
from io import BytesIO

data, fs = sf.read('sound1.wav', dtype='float32')

print(data.dtype)
print(data.shape)

# sd.play(data, fs)
# status = sd.wait()

# Zadanie 1

sf.write('sound_L.wav', data[:, 0], fs)
sf.write('sound_R.wav', data[:, 1], fs)

dataL, fsL = sf.read('sound_L.wav', dtype='float32')
dataR, fsR = sf.read('sound_R.wav', dtype='float32')

mix = dataL + dataR
sf.write('sound_mix.wav', mix, fsL)

# Część trzecia

xaxis = np.arange(0, len(data)) / fs

plt.figure()
plt.title('Dźwięk oryginalny (lewy, prawy, razem)')
plt.subplot(3,1,1)
plt.plot(xaxis, data[:,0])
plt.ylabel('Amplituda')
plt.subplot(3,1,2)
plt.plot(xaxis, data[:,1])
plt.subplot(3,1,3)
plt.plot(xaxis, data[:,0])
plt.plot(xaxis, data[:,1])
plt.xlabel('Czas (s)')
plt.savefig("oryginalny.jpg")
# plt.show()

dataM, fsM = sf.read('sound_mix.wav', dtype='float32')
xaxisM = np.arange(0, len(dataM)) / fsM
plt.figure()
plt.title('Mix - porównanie oryginału z mixem poniżej')
plt.subplot(2,1,1)
plt.plot(xaxis, data)
plt.ylabel('Oryginał')
plt.subplot(2,1,2)
plt.plot(xaxisM, dataM)
plt.ylabel('Mix')
plt.xlabel('Czas (s)')
plt.savefig("mix.jpg")
# plt.show()

# widmo

data, fs = sf.read('sin_440Hz.wav', dtype=np.int32)

plt.figure()
plt.subplot(2, 1, 1)
plt.plot(np.arange(0, data.shape[0])/fs,data)
plt.subplot(2, 1, 2)
yf = scipy.fftpack.fft(data)
plt.plot(np.arange(0, fs, 1.0*fs / (yf.size)), np.abs(yf))
# plt.show()

fsize = 2**8
plt.figure()
plt.subplot(2, 1, 1)
plt.plot(np.arange(0, data.shape[0])/fs, data)
plt.subplot(2, 1, 2)
yf = scipy.fftpack.fft(data, fsize)
plt.plot(np.arange(0, fs/2, fs/fsize), 20*np.log10(np.abs(yf[:fsize//2])))
# plt.show()

# ############

Signal, Fs = sf.read('sin_440Hz.wav', dtype='float32')
def plotAudio(Signal, Fs, TimeMargin):
    xaxis = np.arange(0, len(Signal)) / Fs
    plt.figure()
    plt.figure(figsize=(13, 13))
    plt.title('Sygnał')
    plt.subplot(2, 1, 1)
    # plt.plot(xaxis, Signal)
    plt.plot(np.arange(0, Signal.shape[0]) / Fs, Signal)
    plt.xlabel('Czas (s)')
    plt.ylabel('Amplituda sygnału')
    plt.xlim(TimeMargin[0], TimeMargin[1])

    fsize = 2 ** 8
    plt.subplot(2, 1, 2)
    plt.title('Widmo')
    yf = scipy.fftpack.fft(Signal, fsize)
    plt.plot(np.arange(0, Fs / 2, Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])))
    plt.xlabel('Częstotliwość [Hz]')
    plt.ylabel('Amplituda widma [dB]')
    plt.savefig("z funkcji.jpg")
    # plt.show()

def plotAudio2(Signal,Fs, axs, fsize=2**8, TimeMargin=[0, 0.02]):
    # plt.figure()
    # plt.subplot(2, 1, 1)
    axs[0].plot(np.arange(0, Signal.shape[0]) / Fs, Signal)
    axs[0].set_xlabel("Czas [s]")
    axs[0].set_ylabel("Amplituda")
    axs[0].set_xlim(TimeMargin)
    # plt.subplot(2, 1, 2)
    yf = scipy.fftpack.fft(Signal, fsize)
    x = np.arange(0, Fs / 2, Fs / fsize)
    y = 20 * np.log10(np.abs(yf[:fsize // 2]))
    axs[1].plot(x,y)
    axs[1].set_xlabel("Częstotliwość [Hz]")
    axs[1].set_ylabel("Amplituda [dB]")
    return x[np.argmax(y)],y[np.argmax(y)]
    # plt.show()

# plotAudio(Signal, Fs, TimeMargin=[0, 0.02])

# zadanie 3

document = Document()
document.add_heading('Łukasz Kaszewski - kl49864\nSystemy multimedialne - lab01', 0)  # tworzenie nagłówków druga wartość to poziom nagłówka

files = ['sin_60Hz.wav', 'sin_440Hz.wav', 'sin_8000Hz.wav']
Margins = [2**8, 2**12, 2**16]
for file in files:
    document.add_heading('Plik - {}'.format(file), 2)
    for i, Margin in enumerate(Margins):
        document.add_heading('Parametr fsize {}'.format(Margin), 3)  # nagłówek sekcji, mozę być poziom wyżej
        fig, axs = plt.subplots(2, 1, figsize=(10, 7))  # tworzenie plota

        ############################################################
        # Tu wykonujesz jakieś funkcje i rysujesz wykresy
        Signal, Fs = sf.read(file, dtype='float32')

        # v.1
        # plt.subplot(2, 1, 1)
        # plt.xlim(0, 0.02)
        # plt.plot(np.arange(0, Signal.shape[0]) / Fs, Signal)
        #
        # plt.subplot(2, 1, 2)
        # yf = scipy.fftpack.fft(Signal, Margin)
        # plt.plot(np.arange(0, Fs / 2, Fs / Margin), 20 * np.log10(np.abs(yf[:Margin // 2])))

        # v.2
        result = plotAudio2(Signal, Fs, axs, fsize=Margin)

        ############################################################

        fig.suptitle('fsize {}'.format(Margin))  # Tytuł wykresu
        fig.tight_layout(pad=1.5)  # poprawa czytelności
        memfile = BytesIO()  # tworzenie bufora
        fig.savefig(memfile)  # z zapis do bufora

        document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku

        memfile.close()
        ############################################################
        # Tu dodajesz dane tekstowe - wartosci, wyjscie funkcji ect.
        document.add_paragraph('Wartość najwyższa w miejscu {} Hz o wartości {} dB'.format(result[0], result[1]))
        ############################################################

document.save('kl49864 Łukasz Kaszewski.docx')  # zapis do pliku